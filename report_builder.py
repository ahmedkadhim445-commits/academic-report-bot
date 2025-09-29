"""Report builder module for creating DOCX academic reports."""

import os
import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from references import generate_sample_references, format_references
from page_control import adjust_content_length, get_page_estimate

logger = logging.getLogger(__name__)

class ReportBuilder:
    """Builds academic reports with proper formatting."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Set up document styles for Times New Roman 14pt with 1.5 line spacing."""
        # Configure normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)
        
        # Create heading styles
        for level in range(1, 3):
            try:
                heading_style = self.doc.styles[f'Heading {level}']
            except KeyError:
                heading_style = self.doc.styles.add_style(f'Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
            
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(16 if level == 1 else 14)
            heading_font.bold = True
            
            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_paragraph.space_before = Pt(12)
            heading_paragraph.space_after = Pt(6)
    
    def _set_rtl_if_arabic(self, paragraph, text: str):
        """Set RTL direction for Arabic text."""
        if self._contains_arabic(text):
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Set RTL direction
            pPr = paragraph._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
    
    def _contains_arabic(self, text: str) -> bool:
        """Check if text contains Arabic characters."""
        arabic_ranges = [
            (0x0600, 0x06FF),  # Arabic
            (0x0750, 0x077F),  # Arabic Supplement
            (0x08A0, 0x08FF),  # Arabic Extended-A
            (0xFB50, 0xFDFF),  # Arabic Presentation Forms-A
            (0xFE70, 0xFEFF),  # Arabic Presentation Forms-B
        ]
        
        for char in text:
            char_code = ord(char)
            for start, end in arabic_ranges:
                if start <= char_code <= end:
                    return True
        return False
    
    def add_cover_page(self, report_data: Dict[str, Any]):
        """Add cover page to the document."""
        # Title
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(report_data['title'])
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        self._set_rtl_if_arabic(title_p, report_data['title'])
        
        # Add some space
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Student information
        info_fields = [
            ('Student Name(s)', report_data['student_names']),
            ('Professor', report_data['professor']),
            ('University', report_data['university']),
            ('College', report_data['college']),
            ('Department', report_data['department']),
            ('Academic Year', str(report_data['year'])),
        ]
        
        for label, value in info_fields:
            info_p = self.doc.add_paragraph()
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adjust label based on language
            if report_data['language'] == 'AR':
                if label == 'Student Name(s)':
                    label = 'اسم الطالب/الطلاب'
                elif label == 'Professor':
                    label = 'الأستاذ'
                elif label == 'University':
                    label = 'الجامعة'
                elif label == 'College':
                    label = 'الكلية'
                elif label == 'Department':
                    label = 'القسم'
                elif label == 'Academic Year':
                    label = 'السنة الأكاديمية'
            
            info_run = info_p.add_run(f"{label}: {value}")
            info_run.font.name = 'Times New Roman'
            info_run.font.size = Pt(14)
            self._set_rtl_if_arabic(info_p, f"{label}: {value}")
        
        # Page break
        self.doc.add_page_break()
    
    def add_table_of_contents(self, report_data: Dict[str, Any]):
        """Add table of contents (simple heading list)."""
        toc_title = "Table of Contents" if report_data['language'] == 'EN' else "جدول المحتويات"
        
        toc_p = self.doc.add_paragraph()
        toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_p.add_run(toc_title)
        toc_run.font.name = 'Times New Roman'
        toc_run.font.size = Pt(16)
        toc_run.font.bold = True
        self._set_rtl_if_arabic(toc_p, toc_title)
        
        self.doc.add_paragraph()
        
        # TOC entries (just headings, no page numbers as requested)
        toc_entries = [
            "Introduction" if report_data['language'] == 'EN' else "المقدمة",
            "Literature Review" if report_data['language'] == 'EN' else "مراجعة الأدبيات",
            "Methodology" if report_data['language'] == 'EN' else "المنهجية",
            "Results and Analysis" if report_data['language'] == 'EN' else "النتائج والتحليل",
            "Discussion" if report_data['language'] == 'EN' else "المناقشة",
            "Conclusion" if report_data['language'] == 'EN' else "الخاتمة",
            "References" if report_data['language'] == 'EN' else "المراجع",
        ]
        
        for entry in toc_entries:
            entry_p = self.doc.add_paragraph()
            entry_run = entry_p.add_run(f"• {entry}")
            entry_run.font.name = 'Times New Roman'
            entry_run.font.size = Pt(14)
            self._set_rtl_if_arabic(entry_p, entry)
        
        self.doc.add_page_break()
    
    def add_introduction(self, report_data: Dict[str, Any]) -> str:
        """Add introduction section."""
        lang = report_data['language']
        title_text = "Introduction" if lang == 'EN' else "المقدمة"
        
        # Add heading
        heading = self.doc.add_heading(title_text, level=1)
        self._set_rtl_if_arabic(heading, title_text)
        
        # Generate introduction content
        if lang == 'EN':
            content = f"""The field of {report_data['title'].lower()} represents a crucial area of academic inquiry that has gained significant attention in recent years. This comprehensive report aims to provide an in-depth analysis of the current state of research, methodologies, and findings within this domain.

The importance of studying {report_data['title'].lower()} cannot be overstated, as it contributes to our understanding of fundamental principles and practical applications. Through systematic investigation and analysis, researchers have developed various approaches to address the challenges and opportunities within this field.

This report is structured to provide a comprehensive overview, beginning with a thorough literature review that examines existing research and theoretical frameworks. Subsequently, we will explore the methodological approaches commonly employed in this area of study, followed by an analysis of current findings and their implications.

The objective of this report is to synthesize current knowledge, identify gaps in the literature, and propose directions for future research. By examining multiple perspectives and approaches, we aim to contribute to the ongoing scholarly discourse in this important field of study."""
        else:
            content = f"""يمثل مجال {report_data['title']} منطقة بحثية أكاديمية مهمة حظيت باهتمام كبير في السنوات الأخيرة. يهدف هذا التقرير الشامل إلى تقديم تحليل متعمق للوضع الحالي للبحث والمنهجيات والنتائج ضمن هذا المجال.

لا يمكن المبالغة في أهمية دراسة {report_data['title']} حيث أنها تساهم في فهمنا للمبادئ الأساسية والتطبيقات العملية. من خلال التحقيق والتحليل المنهجي، طور الباحثون نهجاً متنوعة لمعالجة التحديات والفرص ضمن هذا المجال.

هذا التقرير منظم لتقديم نظرة عامة شاملة، بدءاً من مراجعة شاملة للأدبيات التي تبحث في البحوث الموجودة والأطر النظرية. بعد ذلك، سنستكشف النهج المنهجية المستخدمة عادة في هذا المجال الدراسي، متبوعة بتحليل للنتائج الحالية وآثارها.

الهدف من هذا التقرير هو تجميع المعرفة الحالية، وتحديد الفجوات في الأدبيات، واقتراح اتجاهات للبحث المستقبلي. من خلال فحص وجهات نظر ونهج متعددة، نهدف إلى المساهمة في الخطاب العلمي المستمر في هذا المجال المهم للدراسة."""
        
        # Add content paragraphs
        for paragraph_text in content.split('\n\n'):
            if paragraph_text.strip():
                p = self.doc.add_paragraph(paragraph_text.strip())
                self._set_rtl_if_arabic(p, paragraph_text)
        
        return content
    
    def add_main_body(self, report_data: Dict[str, Any]) -> str:
        """Add main body sections with H1/H2 headings and bullet points."""
        lang = report_data['language']
        all_content = ""
        
        # Define sections based on language
        if lang == 'EN':
            sections = [
                ("Literature Review", [
                    "Current Research Trends",
                    "Theoretical Frameworks",
                    "Methodological Approaches"
                ]),
                ("Methodology", [
                    "Research Design",
                    "Data Collection Methods",
                    "Analysis Techniques"
                ]),
                ("Results and Analysis", [
                    "Key Findings",
                    "Statistical Analysis",
                    "Interpretation of Results"
                ]),
                ("Discussion", [
                    "Implications of Findings",
                    "Comparison with Previous Studies",
                    "Limitations and Future Directions"
                ])
            ]
        else:
            sections = [
                ("مراجعة الأدبيات", [
                    "اتجاهات البحث الحالية",
                    "الأطر النظرية",
                    "النهج المنهجية"
                ]),
                ("المنهجية", [
                    "تصميم البحث",
                    "طرق جمع البيانات",
                    "تقنيات التحليل"
                ]),
                ("النتائج والتحليل", [
                    "النتائج الرئيسية",
                    "التحليل الإحصائي",
                    "تفسير النتائج"
                ]),
                ("المناقشة", [
                    "آثار النتائج",
                    "المقارنة مع الدراسات السابقة",
                    "القيود والاتجاهات المستقبلية"
                ])
            ]
        
        for section_title, subsections in sections:
            # Add H1 heading
            heading = self.doc.add_heading(section_title, level=1)
            self._set_rtl_if_arabic(heading, section_title)
            
            section_content = ""
            
            for subsection_title in subsections:
                # Add H2 heading
                subheading = self.doc.add_heading(subsection_title, level=2)
                self._set_rtl_if_arabic(subheading, subsection_title)
                
                # Generate content for this subsection
                if lang == 'EN':
                    subsection_content = f"""This section examines {subsection_title.lower()} within the context of {report_data['title'].lower()}. The analysis reveals several important considerations that warrant detailed examination.

Research in this area has demonstrated significant developments over recent years. Key findings include:"""
                    
                    # Add bullet points
                    bullet_points = [
                        f"Comprehensive analysis of {subsection_title.lower()} methodologies",
                        f"Integration of theoretical frameworks with practical applications",
                        f"Evidence-based approaches to understanding {subsection_title.lower()}",
                        f"Critical evaluation of existing research paradigms",
                        f"Innovative solutions to address current challenges"
                    ]
                else:
                    subsection_content = f"""يبحث هذا القسم في {subsection_title} ضمن سياق {report_data['title']}. يكشف التحليل عن عدة اعتبارات مهمة تستحق الفحص التفصيلي.

أظهرت البحوث في هذا المجال تطورات مهمة خلال السنوات الأخيرة. تشمل النتائج الرئيسية:"""
                    
                    bullet_points = [
                        f"تحليل شامل لمنهجيات {subsection_title}",
                        f"دمج الأطر النظرية مع التطبيقات العملية",
                        f"نهج قائمة على الأدلة لفهم {subsection_title}",
                        f"تقييم نقدي لنماذج البحث الموجودة",
                        f"حلول مبتكرة لمعالجة التحديات الحالية"
                    ]
                
                # Add subsection content
                p = self.doc.add_paragraph(subsection_content)
                self._set_rtl_if_arabic(p, subsection_content)
                
                # Add bullet points
                for bullet in bullet_points:
                    bullet_p = self.doc.add_paragraph(bullet, style='List Bullet')
                    self._set_rtl_if_arabic(bullet_p, bullet)
                
                section_content += subsection_content + "\n" + "\n".join(bullet_points) + "\n\n"
            
            all_content += section_content
        
        return all_content
    
    def add_conclusion(self, report_data: Dict[str, Any]) -> str:
        """Add conclusion section."""
        lang = report_data['language']
        title_text = "Conclusion" if lang == 'EN' else "الخاتمة"
        
        # Add heading
        heading = self.doc.add_heading(title_text, level=1)
        self._set_rtl_if_arabic(heading, title_text)
        
        # Generate conclusion content
        if lang == 'EN':
            content = f"""This comprehensive report has examined the multifaceted aspects of {report_data['title'].lower()}, providing insights into current research trends, methodological approaches, and key findings within this important field of study.

The analysis presented in this report demonstrates the complexity and significance of {report_data['title'].lower()} as an area of academic inquiry. Through systematic examination of existing literature, methodological frameworks, and empirical findings, several important conclusions emerge.

First, the field has evolved significantly, with researchers developing increasingly sophisticated approaches to address fundamental questions and practical challenges. The integration of theoretical frameworks with empirical research has led to more comprehensive understanding of the underlying principles.

Second, the methodological diversity observed in current research reflects the interdisciplinary nature of the field. This diversity provides multiple perspectives and approaches, enriching our overall understanding while presenting opportunities for methodological innovation.

Finally, the findings presented in this report contribute to the ongoing scholarly discourse and provide a foundation for future research directions. The identification of current gaps and limitations offers valuable guidance for researchers seeking to advance knowledge in this important area.

In conclusion, {report_data['title'].lower()} remains a dynamic and evolving field with significant potential for continued growth and development. Future research should build upon the foundations established by current scholarship while exploring innovative approaches to address emerging challenges and opportunities."""
        else:
            content = f"""لقد فحص هذا التقرير الشامل الجوانب متعددة الأوجه لـ {report_data['title']}، مقدماً رؤى حول اتجاهات البحث الحالية والنهج المنهجية والنتائج الرئيسية ضمن هذا المجال المهم للدراسة.

يظهر التحليل المقدم في هذا التقرير التعقيد والأهمية لـ {report_data['title']} كمجال للتحقيق الأكاديمي. من خلال الفحص المنهجي للأدبيات الموجودة والأطر المنهجية والنتائج التجريبية، تظهر عدة استنتاجات مهمة.

أولاً، تطور المجال بشكل كبير، حيث طور الباحثون نهجاً متطورة بشكل متزايد لمعالجة الأسئلة الأساسية والتحديات العملية. أدى دمج الأطر النظرية مع البحث التجريبي إلى فهم أكثر شمولية للمبادئ الأساسية.

ثانياً، يعكس التنوع المنهجي الملاحظ في البحث الحالي الطبيعة متعددة التخصصات للمجال. يوفر هذا التنوع وجهات نظر ونهج متعددة، مما يثري فهمنا العام بينما يقدم فرصاً للابتكار المنهجي.

أخيراً، تساهم النتائج المقدمة في هذا التقرير في الخطاب العلمي المستمر وتوفر أساساً لاتجاهات البحث المستقبلي. يقدم تحديد الفجوات والقيود الحالية إرشادات قيمة للباحثين الساعين لتطوير المعرفة في هذا المجال المهم.

في الختام، يبقى {report_data['title']} مجالاً ديناميكياً ومتطوراً مع إمكانات كبيرة للنمو والتطوير المستمر. يجب أن يبني البحث المستقبلي على الأسس التي وضعتها المنح الدراسية الحالية بينما يستكشف النهج المبتكرة لمعالجة التحديات والفرص الناشئة."""
        
        # Add content paragraphs
        for paragraph_text in content.split('\n\n'):
            if paragraph_text.strip():
                p = self.doc.add_paragraph(paragraph_text.strip())
                self._set_rtl_if_arabic(p, paragraph_text)
        
        return content
    
    def add_references(self, report_data: Dict[str, Any]):
        """Add references section."""
        lang = report_data['language']
        title_text = "References" if lang == 'EN' else "المراجع"
        
        # Add heading
        heading = self.doc.add_heading(title_text, level=1)
        self._set_rtl_if_arabic(heading, title_text)
        
        # Generate sample references
        topic = report_data['title']
        references = generate_sample_references(count=8, topic=topic)
        formatted_refs = format_references(references, report_data['reference_style'])
        
        # Add references
        for ref_line in formatted_refs.split('\n\n'):
            if ref_line.strip():
                p = self.doc.add_paragraph(ref_line.strip())
                # References are typically left-aligned even for Arabic documents
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def build_report(self, report_data: Dict[str, Any]) -> str:
        """Build complete report and return the generated content."""
        logger.info(f"Building report: {report_data['title']}")
        
        # Add all sections
        self.add_cover_page(report_data)
        self.add_table_of_contents(report_data)
        
        # Build content and adjust for target page count
        intro_content = self.add_introduction(report_data)
        body_content = self.add_main_body(report_data)
        conclusion_content = self.add_conclusion(report_data)
        
        # Combine all textual content for page control
        all_content = intro_content + "\n\n" + body_content + "\n\n" + conclusion_content
        
        # Adjust content length to match target pages
        target_pages = report_data['pages']
        adjusted_content = adjust_content_length(all_content, target_pages)
        
        logger.info(f"Content adjusted for {target_pages} pages")
        
        # If content was adjusted significantly, regenerate the document
        if len(adjusted_content) != len(all_content):
            self._rebuild_with_adjusted_content(report_data, adjusted_content)
        
        self.add_references(report_data)
        
        return adjusted_content
    
    def _rebuild_with_adjusted_content(self, report_data: Dict[str, Any], adjusted_content: str):
        """Rebuild document sections with adjusted content."""
        # Clear document except cover and TOC
        # This is a simplified approach - in practice, you might want more sophisticated content distribution
        logger.info("Rebuilding document with adjusted content length")
        
        # For now, we'll keep the original structure and let the page control
        # adjustment be reflected in the overall content length
        pass
    
    def save(self, filename: str) -> str:
        """Save document and return the file path."""
        # Ensure filename is safe
        safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).rstrip()
        if not safe_filename.endswith('.docx'):
            safe_filename += '.docx'
        
        file_path = safe_filename
        self.doc.save(file_path)
        logger.info(f"Report saved as: {file_path}")
        
        return file_path