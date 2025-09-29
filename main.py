import os, io, re, tempfile
from datetime import datetime
from dotenv import load_dotenv

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ConversationHandler, ContextTypes, filters
)

# DOCX
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF (اختياري)
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

load_dotenv()
BOT_TOKEN = os.getenv("BOT_TOKEN")

# محادثة
(TITLE, LANG, STUDENT, PROFESSOR, UNIVERSITY, COLLEGE, DEPARTMENT,
 YEAR, PAGES, REFSTYLE, CONFIRM) = range(11)

# أنماط المراجع المدعومة
REF_STYLES = ["APA", "IEEE", "MLA", "Harvard", "Chicago"]

def _slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9_\-]+", "_", s).strip("_")

def _set_paragraph_style(p):
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)

def build_docx(data: dict) -> bytes:
    doc = Document()
    # هوامش
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # الغلاف
    doc.add_paragraph()
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = head.add_run(data["university"] or "University")
    r.bold = True; r.font.size = Pt(16)

    head.add_run("\n")
    r = head.add_run(data["college"] or "College"); r.bold = True; r.font.size = Pt(14)
    head.add_run("\n")
    r = head.add_run(data["department"] or "Department"); r.bold = True; r.font.size = Pt(13)

    doc.add_paragraph("\n\n")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(f"\"{data['title']}\""); r.bold = True; r.font.size = Pt(22)

    doc.add_paragraph("\n\n")

    info = doc.add_paragraph()
    info.add_run(f"Student(s): {data['student']}\n")
    info.add_run(f"Professor: {data['professor']}\n")
    info.add_run(f"Stage/Year: {data['year']}\n")
    info.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
    for r in info.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    doc.add_page_break()

    # فهرس (نصي مبسّط – المستخدم يقدر يولد TOC تلقائي من Word لاحقاً)
    h = doc.add_heading("TABLE OF CONTENTS", level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("2. Main Body")
    doc.add_paragraph("3. Conclusion")
    doc.add_paragraph("4. References")
    doc.add_page_break()

    # مقدمة
    h = doc.add_heading("Introduction", level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = (
        f"This report titled \"{data['title']}\" explores the topic in a structured academic style. "
        f"It belongs to {data['department']} / {data['college']} at {data['university']}."
        if data["language"] == "english" else
        f"يتناول هذا التقرير بعنوان \"{data['title']}\" الموضوع بصورة أكاديمية منظمة، "
        f"وهو تابع لـ {data['department']} / {data['college']} في {data['university']}."
    )
    p = doc.add_paragraph(intro); _set_paragraph_style(p)

    # متن
    doc.add_page_break()
    h = doc.add_heading("Main Body", level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bullets = [
        ("Background & Importance" if data["language"]=="english" else "الخلفية والأهمية"),
        ("Key Concepts & Definitions" if data["language"]=="english" else "المفاهيم والتعاريف"),
        ("Methods / Approaches" if data["language"]=="english" else "المنهجيات/الأساليب"),
        ("Applications / Case Studies" if data["language"]=="english" else "التطبيقات/دراسات الحالة"),
        ("Challenges & Future Work" if data["language"]=="english" else "التحديات والاتجاهات المستقبلية"),
    ]
    for b in bullets:
        p = doc.add_paragraph(b); p.runs[0].bold = True; _set_paragraph_style(p)
        p = doc.add_paragraph(
            "• " + (
                "Write a concise, well-argued paragraph for this point."
                if data["language"]=="english"
                else "اكتب فقرة موجزة ومتماسكة لهذه النقطة."
            )
        ); _set_paragraph_style(p)

    # خاتمة
    doc.add_page_break()
    h = doc.add_heading("Conclusion", level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    concl = (
        "This report summarized the main ideas, practical implications, and suggested directions for future work."
        if data["language"]=="english" else
        "قدّم هذا التقرير خلاصةً للأفكار الأساسية والانعكاسات العملية مع مقترحات للعمل المستقبلي."
    )
    p = doc.add_paragraph(concl); _set_paragraph_style(p)

    # مراجع
    doc.add_page_break()
    h = doc.add_heading("References", level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    refs = [
        {"author": "Doe, J.", "year": "2022", "title": "Sample Paper", "source": "Journal of Examples"},
        {"author": "Smith, A.", "year": "2021", "title": "Another Work", "source": "Conference on Samples"},
    ]
    for i, ref in enumerate(refs, 1):
        if data["refstyle"] == "IEEE":
            line = f"[{i}] {ref['author']}, \"{ref['title']}\", {ref['source']}, {ref['year']}."
        elif data["refstyle"] == "MLA":
            line = f"{ref['author']} \"{ref['title']}\". {ref['source']}, {ref['year']}."
        elif data["refstyle"] == "Harvard":
            line = f"{ref['author']} ({ref['year']}). {ref['title']}. {ref['source']}."
        elif data["refstyle"] == "Chicago":
            line = f"{ref['author']}. {ref['year']}. {ref['title']}. {ref['source']}."
        else:  # APA
            line = f"{ref['author']} ({ref['year']}). {ref['title']}. {ref['source']}."
        p = doc.add_paragraph(line); _set_paragraph_style(p)

    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.read()

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("أهلين! ارسل عنوان التقرير (Report Title).")
    return TITLE

async def title_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["title"] = update.message.text.strip()
    kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("🇮🇶 العربية", callback_data="lang_ar")],
        [InlineKeyboardButton("🇬🇧 English", callback_data="lang_en")],
    ])
    await update.message.reply_text("اختار لغة التقرير:", reply_markup=kb)
    return LANG

async def lang_cb(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query; await q.answer()
    context.user_data["language"] = "arabic" if q.data=="lang_ar" else "english"
    await q.message.reply_text("اكتب اسم الطالب (أو أسماء الطلاب).")
    return STUDENT

async def student_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["student"] = update.message.text.strip()
    await update.message.reply_text("اكتب اسم الدكتور/الأستاذ.")
    return PROFESSOR

async def professor_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["professor"] = update.message.text.strip()
    await update.message.reply_text("اكتب اسم الجامعة.")
    return UNIVERSITY

async def university_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["university"] = update.message.text.strip()
    await update.message.reply_text("اكتب اسم الكلية.")
    return COLLEGE

async def college_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["college"] = update.message.text.strip()
    await update.message.reply_text("اكتب اسم القسم.")
    return DEPARTMENT

async def department_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["department"] = update.message.text.strip()
    await update.message.reply_text("اكتب المرحلة/السنة الدراسية.")
    return YEAR

async def year_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["year"] = update.message.text.strip()
    await update.message.reply_text("كم صفحة تريد؟ (بين 5 و 40)")
    return PAGES

async def pages_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        pages = int(update.message.text.strip())
        if pages < 5 or pages > 40:
            await update.message.reply_text("الرجاء رقم بين 5 و 40.")
            return PAGES
        context.user_data["pages"] = pages
    except:
        await update.message.reply_text("الرجاء إرسال رقم صحيح.")
        return PAGES

    kb = InlineKeyboardMarkup([[InlineKeyboardButton(s, callback_data=f"ref_{s}") ] for s in REF_STYLES])
    await update.message.reply_text("اختر نمط المراجع:", reply_markup=kb)
    return REFSTYLE

async def ref_cb(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query; await q.answer()
    context.user_data["refstyle"] = q.data.replace("ref_","")
    data = context.user_data
    txt = (
        f"معاينة البيانات:\n"
        f"العنوان: {data['title']}\nاللغة: {data['language']}\nالطالب: {data['student']}\n"
        f"الأستاذ: {data['professor']}\nالجامعة: {data['university']}\n"
        f"الكلية: {data['college']}\nالقسم: {data['department']}\n"
        f"السنة/المرحلة: {data['year']}\nالصفحات: {data['pages']}\n"
        f"نمط المراجع: {data['refstyle']}\n\nتأكيد الإنشاء؟"
    )
    kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("✅ نعم", callback_data="go")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="cancel")],
    ])
    await q.message.reply_text(txt, reply_markup=kb)
    return CONFIRM

async def confirm_cb(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query; await q.answer()
    if q.data == "cancel":
        await q.message.reply_text("تم الإلغاء.")
        return ConversationHandler.END

    data = context.user_data
    # توليد DOCX
    docx_bytes = build_docx(data)

    # اسم ملف
    base = _slug(f"{data['title']}")
    docx_name = f"{base}.docx"
    pdf_name  = f"{base}.pdf"

    # أرسل DOCX
    await q.message.reply_document(document=io.BytesIO(docx_bytes), filename=docx_name,
                                   caption="تم إنشاء تقرير Word ✅")

    # حاول PDF إذا متاح
    if DOCX2PDF_AVAILABLE:
        with tempfile.TemporaryDirectory() as td:
            path_docx = os.path.join(td, docx_name)
            with open(path_docx, "wb") as f: f.write(docx_bytes)
            try:
                docx2pdf_convert(path_docx, os.path.join(td, pdf_name))
                with open(os.path.join(td, pdf_name), "rb") as f:
                    await q.message.reply_document(document=io.BytesIO(f.read()), filename=pdf_name,
                                                   caption="نسخة PDF ✅")
            except Exception:
                await q.message.reply_text("لم يتمكن البوت من توليد PDF على هذا الخادم. أرسلنا ملف Word فقط.")
    else:
        await q.message.reply_text("تحويل PDF غير مفعّل على هذا النظام. تم إرسال ملف Word فقط.")

    await q.message.reply_text("للبدء من جديد أرسل /start")
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("ألغينا العملية. أرسل /start للبدء من جديد.")
    return ConversationHandler.END

def main():
    if not BOT_TOKEN:
        raise RuntimeError("BOT_TOKEN غير موجود. ضعه كمتغير بيئة.")

    app = Application.builder().token(BOT_TOKEN).build()

    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, title_step)],
            LANG:  [CallbackQueryHandler(lang_cb, pattern="^lang_")],
            STUDENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, student_step)],
            PROFESSOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, professor_step)],
            UNIVERSITY: [MessageHandler(filters.TEXT & ~filters.COMMAND, university_step)],
            COLLEGE: [MessageHandler(filters.TEXT & ~filters.COMMAND, college_step)],
            DEPARTMENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, department_step)],
            YEAR: [MessageHandler(filters.TEXT & ~filters.COMMAND, year_step)],
            PAGES: [MessageHandler(filters.TEXT & ~filters.COMMAND, pages_step)],
            REFSTYLE: [CallbackQueryHandler(ref_cb, pattern="^ref_")],
            CONFIRM: [CallbackQueryHandler(confirm_cb, pattern="^(go|cancel)$")],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
        per_chat=True,
    )

    app.add_handler(conv)
    app.run_polling()

if __name__ == "__main__":
    main()
