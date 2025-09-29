"""
Academic Report Generator - Creates professional academic reports in Word and PDF formats
"""

import os
import tempfile
import subprocess
from datetime import datetime
from typing import Dict, Any, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class ReportGenerator:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        
    def _set_font_and_spacing(self, paragraph, font_name="Times New Roman", font_size=14):
        """Set font and line spacing for a paragraph."""
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.5
        
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
    
    def _add_styled_paragraph(self, doc, text, style_name=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """Add a paragraph with consistent styling."""
        if style_name:
            paragraph = doc.add_paragraph(text, style=style_name)
        else:
            paragraph = doc.add_paragraph(text)
        
        paragraph.alignment = alignment
        self._set_font_and_spacing(paragraph)
        return paragraph
    
    def _create_cover_page(self, doc: Document, data: Dict[str, Any]):
        """Create the cover page."""
        # University name (centered, large)
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(data['university'])
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(18)
        title_run.bold = True
        
        # College name
        college_p = doc.add_paragraph()
        college_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        college_run = college_p.add_run(data['college'])
        college_run.font.name = "Times New Roman"
        college_run.font.size = Pt(16)
        college_run.bold = True
        
        # Department name
        dept_p = doc.add_paragraph()
        dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_run = dept_p.add_run(data['department'])
        dept_run.font.name = "Times New Roman"
        dept_run.font.size = Pt(14)
        
        # Add some space
        for _ in range(3):
            doc.add_paragraph()
        
        # Report title (centered, large, bold)
        report_title_p = doc.add_paragraph()
        report_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title_run = report_title_p.add_run(data['title'])
        report_title_run.font.name = "Times New Roman"
        report_title_run.font.size = Pt(20)
        report_title_run.bold = True
        
        # Add space
        for _ in range(5):
            doc.add_paragraph()
        
        # Student info (right aligned)
        student_info = [
            f"Student: {data['student']}",
            f"Professor: {data['professor']}",
            f"Academic Year: {data['year']}"
        ]
        
        for info in student_info:
            info_p = doc.add_paragraph()
            info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            info_run = info_p.add_run(info)
            info_run.font.name = "Times New Roman"
            info_run.font.size = Pt(14)
        
        # Page break
        doc.add_page_break()
    
    def _create_table_of_contents(self, doc: Document, data: Dict[str, Any]):
        """Create table of contents."""
        if data['language'] == 'AR':
            toc_title = 'جدول المحتويات'
            toc_items = [
                "1. المقدمة",
                "2. مراجعة الأدبيات", 
                "3. المنهجية",
                "4. النتائج والمناقشة",
                "5. التحليل",
                "6. الخاتمة",
                "7. المراجع"
            ]
        else:
            toc_title = 'Table of Contents'
            toc_items = [
                "1. Introduction",
                "2. Literature Review", 
                "3. Methodology",
                "4. Results and Discussion",
                "5. Analysis",
                "6. Conclusion",
                "7. References"
            ]
        
        toc_heading = doc.add_heading(toc_title, level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font_and_spacing(toc_heading, font_size=16)
        
        for item in toc_items:
            toc_p = doc.add_paragraph(item)
            toc_p.paragraph_format.left_indent = Inches(0.5)
            # Align Arabic text to the right
            if data['language'] == 'AR':
                toc_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._set_font_and_spacing(toc_p)
        
        doc.add_page_break()
    
    def _create_introduction(self, doc: Document, data: Dict[str, Any]):
        """Create introduction section."""
        if data['language'] == 'AR':
            heading_text = '1. المقدمة'
            intro_text = f"""
            يقدم هذا التقرير تحليلاً شاملاً لموضوع {data['title']}. تهدف هذه الدراسة إلى توفير رؤى قيمة والمساهمة في المعرفة الموجودة في هذا المجال.
            
            الهدف الأساسي من هذا البحث هو فحص الجوانب المختلفة للموضوع وتقديم النتائج التي تتسم بالصرامة الأكاديمية والأهمية العملية. تحدد هذه المقدمة نطاق الدراسة ومنهجيتها وأهميتها.
            
            يتم تنظيم التقرير لتوفير تدفق منطقي للمعلومات، بدءاً من مراجعة شاملة للأدبيات، متبوعة بالمنهجية المستخدمة، وعرض النتائج، والتحليل المفصل، والملاحظات الختامية.
            """
        else:
            heading_text = '1. Introduction'
            intro_text = f"""
            This report presents a comprehensive analysis of {data['title']}. The study aims to provide valuable insights and contribute to the existing body of knowledge in this field. 
            
            The primary objective of this research is to examine the various aspects of the topic and present findings that are both academically rigorous and practically relevant. This introduction outlines the scope, methodology, and significance of the study.
            
            The report is structured to provide a logical flow of information, beginning with a thorough literature review, followed by the methodology employed, presentation of results, detailed analysis, and concluding remarks. Each section builds upon the previous one to create a cohesive and comprehensive academic document.
            """
        
        heading = doc.add_heading(heading_text, level=1)
        self._set_font_and_spacing(heading, font_size=16)
        
        self._add_styled_paragraph(doc, intro_text.strip())
    
    def _create_body_sections(self, doc: Document, pages: int):
        """Create body sections based on the number of pages."""
        sections = [
            ("2. Literature Review", "This section reviews the existing literature related to the research topic. A comprehensive analysis of previous studies, theories, and findings provides the foundation for this research."),
            ("3. Methodology", "This section describes the research methodology employed in this study. The approach, data collection methods, and analytical techniques are outlined to ensure transparency and reproducibility."),
            ("4. Results and Discussion", "This section presents the findings of the research and discusses their implications. The results are analyzed in the context of the research objectives and existing literature."),
            ("5. Analysis", "This section provides a detailed analysis of the findings, examining patterns, relationships, and significance of the results in relation to the research questions.")
        ]
        
        # Calculate approximate content per section based on pages
        base_paragraphs = max(2, pages // 4)
        
        for title, intro in sections:
            heading = doc.add_heading(title, level=1)
            self._set_font_and_spacing(heading, font_size=16)
            
            # Add introduction paragraph
            self._add_styled_paragraph(doc, intro)
            
            # Add additional content paragraphs
            for i in range(base_paragraphs):
                content = f"""
                This paragraph provides additional analysis and discussion relevant to {title.split('.')[1].strip()}. 
                The content here demonstrates the depth of research and analysis undertaken in this study. 
                Various aspects of the topic are examined to provide comprehensive coverage and insight.
                
                Further elaboration on the key points helps to establish the academic rigor of this work. 
                The discussion incorporates relevant theoretical frameworks and empirical evidence to support the arguments presented.
                """
                self._add_styled_paragraph(doc, content.strip())
    
    def _create_conclusion(self, doc: Document, data: Dict[str, Any]):
        """Create conclusion section."""
        heading = doc.add_heading('6. Conclusion', level=1)
        self._set_font_and_spacing(heading, font_size=16)
        
        conclusion_text = f"""
        This study has provided a comprehensive examination of {data['title']}. The research has successfully addressed the primary objectives and contributed valuable insights to the field.
        
        The findings demonstrate the importance of continued research in this area and highlight several key implications for both theoretical understanding and practical application. The methodology employed has proven effective in generating reliable and valid results.
        
        Future research directions may include expanding the scope of the study, incorporating additional variables, or applying the findings to different contexts. The conclusions drawn from this research provide a solid foundation for such future endeavors.
        
        In summary, this report has successfully explored the research topic through rigorous academic inquiry, providing meaningful contributions to the existing body of knowledge.
        """
        
        self._add_styled_paragraph(doc, conclusion_text.strip())
    
    def _create_references(self, doc: Document, ref_style: str):
        """Create references section."""
        heading = doc.add_heading('7. References', level=1)
        self._set_font_and_spacing(heading, font_size=16)
        
        # Sample references in different styles
        references = {
            'APA': [
                "Smith, J. A. (2023). Academic Research Methods in Modern Education. Journal of Educational Studies, 45(3), 123-145.",
                "Johnson, M. B., & Davis, K. L. (2022). Contemporary Approaches to Academic Writing. Academic Press.",
                "Brown, R. C. (2024). Research Methodology and Data Analysis. International Journal of Research, 12(2), 67-89.",
                "Wilson, A. D., Thompson, E. F., & Lee, S. M. (2023). Advanced Topics in Academic Research. University Publications.",
                "Garcia, L. P. (2022). Statistical Analysis for Academic Research. Research Methods Quarterly, 8(4), 201-218."
            ],
            'IEEE': [
                "J. A. Smith, \"Academic Research Methods in Modern Education,\" Journal of Educational Studies, vol. 45, no. 3, pp. 123-145, 2023.",
                "M. B. Johnson and K. L. Davis, Contemporary Approaches to Academic Writing. Academic Press, 2022.",
                "R. C. Brown, \"Research Methodology and Data Analysis,\" International Journal of Research, vol. 12, no. 2, pp. 67-89, 2024.",
                "A. D. Wilson, E. F. Thompson, and S. M. Lee, Advanced Topics in Academic Research. University Publications, 2023.",
                "L. P. Garcia, \"Statistical Analysis for Academic Research,\" Research Methods Quarterly, vol. 8, no. 4, pp. 201-218, 2022."
            ],
            'MLA': [
                "Smith, John A. \"Academic Research Methods in Modern Education.\" Journal of Educational Studies, vol. 45, no. 3, 2023, pp. 123-145.",
                "Johnson, Mary B., and Karen L. Davis. Contemporary Approaches to Academic Writing. Academic Press, 2022.",
                "Brown, Robert C. \"Research Methodology and Data Analysis.\" International Journal of Research, vol. 12, no. 2, 2024, pp. 67-89.",
                "Wilson, Andrew D., et al. Advanced Topics in Academic Research. University Publications, 2023.",
                "Garcia, Luis P. \"Statistical Analysis for Academic Research.\" Research Methods Quarterly, vol. 8, no. 4, 2022, pp. 201-218."
            ],
            'Harvard': [
                "Smith, J.A., 2023. Academic Research Methods in Modern Education. Journal of Educational Studies, 45(3), pp.123-145.",
                "Johnson, M.B. and Davis, K.L., 2022. Contemporary Approaches to Academic Writing. Academic Press.",
                "Brown, R.C., 2024. Research Methodology and Data Analysis. International Journal of Research, 12(2), pp.67-89.",
                "Wilson, A.D., Thompson, E.F. and Lee, S.M., 2023. Advanced Topics in Academic Research. University Publications.",
                "Garcia, L.P., 2022. Statistical Analysis for Academic Research. Research Methods Quarterly, 8(4), pp.201-218."
            ],
            'Chicago': [
                "Smith, John A. \"Academic Research Methods in Modern Education.\" Journal of Educational Studies 45, no. 3 (2023): 123-145.",
                "Johnson, Mary B., and Karen L. Davis. Contemporary Approaches to Academic Writing. Academic Press, 2022.",
                "Brown, Robert C. \"Research Methodology and Data Analysis.\" International Journal of Research 12, no. 2 (2024): 67-89.",
                "Wilson, Andrew D., Emily F. Thompson, and Susan M. Lee. Advanced Topics in Academic Research. University Publications, 2023.",
                "Garcia, Luis P. \"Statistical Analysis for Academic Research.\" Research Methods Quarterly 8, no. 4 (2022): 201-218."
            ]
        }
        
        ref_list = references.get(ref_style, references['APA'])
        
        for ref in ref_list:
            ref_p = doc.add_paragraph(ref)
            ref_p.paragraph_format.left_indent = Inches(0.5)
            ref_p.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
            self._set_font_and_spacing(ref_p)
    
    async def generate_report(self, data: Dict[str, Any]) -> Tuple[str, str]:
        """Generate academic report in both DOCX and PDF formats."""
        # Create document
        doc = Document()
        
        # Set default font for the document
        styles = doc.styles
        style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Create sections
        self._create_cover_page(doc, data)
        self._create_table_of_contents(doc, data)
        self._create_introduction(doc, data)
        self._create_body_sections(doc, data['pages'])
        self._create_conclusion(doc, data)
        self._create_references(doc, data['ref_style'])
        
        # Save DOCX file
        docx_filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        docx_path = os.path.join(self.temp_dir, docx_filename)
        doc.save(docx_path)
        
        # Convert to PDF
        pdf_filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(self.temp_dir, pdf_filename)
        
        try:
            # Use LibreOffice for PDF conversion (more reliable than docx2pdf)
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', self.temp_dir, docx_path
            ], check=True, capture_output=True)
            
            # Find the generated PDF file
            base_name = os.path.splitext(docx_filename)[0]
            generated_pdf = os.path.join(self.temp_dir, f"{base_name}.pdf")
            
            if os.path.exists(generated_pdf):
                os.rename(generated_pdf, pdf_path)
            else:
                raise Exception("PDF conversion failed")
                
        except Exception as e:
            # Fallback: create a simple PDF with reportlab
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title = Paragraph(f"<b>{data['title']}</b>", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 0.5*inch))
            
            # Add basic info
            info_text = f"""
            <b>Student:</b> {data['student']}<br/>
            <b>Professor:</b> {data['professor']}<br/>
            <b>University:</b> {data['university']}<br/>
            <b>College:</b> {data['college']}<br/>
            <b>Department:</b> {data['department']}<br/>
            <b>Year:</b> {data['year']}<br/>
            <b>Reference Style:</b> {data['ref_style']}
            """
            info_para = Paragraph(info_text, styles['Normal'])
            story.append(info_para)
            
            pdf_doc.build(story)
        
        return docx_path, pdf_path